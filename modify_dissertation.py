#!/usr/bin/env python3
"""
Script to modify dissertation.v5.docx and save as dissertation.v6.docx
Makes 14 specific text changes as specified.
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as etree
import sys

INPUT_PATH = "/Users/jamie/Documents/University/ImpactTracker/dissertation.v5.docx"
OUTPUT_PATH = "/Users/jamie/Documents/University/ImpactTracker/dissertation.v6.docx"


def set_para_text(para, new_text):
    """Clear all runs and set first run to new_text, preserving style."""
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def insert_paragraph_after(ref_para, text, style_name, doc):
    """Insert a new paragraph with given text and style immediately after ref_para."""
    new_p = deepcopy(ref_para._element)
    ref_para._element.addnext(new_p)
    # Find the newly inserted paragraph in the document
    # It should be directly after ref_para in the XML
    # We need to get the Python paragraph object for the new element
    # The easiest way is to search doc.paragraphs for the new element
    for para in doc.paragraphs:
        if para._element is new_p:
            # Clear runs and set text
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            para.style = doc.styles[style_name]
            return para
    return None


def find_para_by_substring(doc, substring):
    """Find first paragraph containing the given substring."""
    for i, para in enumerate(doc.paragraphs):
        if substring in para.text:
            return i, para
    return None, None


def main():
    print(f"Loading {INPUT_PATH}...")
    doc = Document(INPUT_PATH)
    paragraphs = doc.paragraphs
    print(f"Total paragraphs: {len(paragraphs)}")
    print()

    changes_made = 0

    # ------------------------------------------------------------------
    # CHANGE 1 - Para 26 (Section 1.1)
    # ------------------------------------------------------------------
    old_text_1 = "E-commerce reduces friction (recommendations, one-click checkout, fast delivery), which can increase discretionary purchasing."
    new_text_1 = ("E-commerce reduces friction \u2014 recommendations, one-click checkout, fast delivery \u2014 which demonstrably increases "
                  "discretionary purchasing. Global retail e-commerce revenue exceeded $5.8 trillion in 2023 and is forecast to surpass "
                  "$8 trillion by 2027 (Statista, 2024), meaning the environmental cost of each purchase is multiplied at a scale that "
                  "dwarfs individual choice. That cost is distributed across the full product lifecycle: raw material extraction, "
                  "manufacturing, packaging, international shipping, UK distribution, and end-of-life disposal. DSP Eco Tracker makes "
                  "those costs visible at the point of decision, before a purchase is committed, by returning a fast, interpretable "
                  "eco-grade for any Amazon listing.")
    idx, para = find_para_by_substring(doc, old_text_1)
    if para is not None:
        set_para_text(para, new_text_1)
        print(f"CHANGE 1 applied: Para {idx} (Section 1.1) replaced.")
        changes_made += 1
    else:
        print("CHANGE 1 FAILED: Could not find paragraph containing the target text.")

    # ------------------------------------------------------------------
    # CHANGE 2 - Para 28 (Section 1.2)
    # ------------------------------------------------------------------
    old_text_2 = "Most shoppers do not see credible lifecycle information at decision time"
    new_text_2 = ("Most shoppers have no access to credible lifecycle information at decision time. Sustainability labels, where they "
                  "exist at all, are voluntary, inconsistent across platforms, and rarely cover the full supply chain (Th\u00f8gersen et al., "
                  "2010). Full life cycle assessment data is almost never available for individual marketplace listings. This creates a "
                  "well-documented gap: research consistently finds that pro-environmental intention fails to translate into greener "
                  "purchasing when the information required to act on that intention is missing or untrustworthy (Kollmuss and Agyeman, "
                  "2002). A practical intervention is point-of-purchase eco-feedback that is fast enough to fit within a shopping "
                  "decision, honest about uncertainty, and embedded within the platform where the choice is made \u2014 rather than "
                  "requiring users to seek information elsewhere.")
    idx, para = find_para_by_substring(doc, old_text_2)
    if para is not None:
        set_para_text(para, new_text_2)
        print(f"CHANGE 2 applied: Para {idx} (Section 1.2) replaced.")
        changes_made += 1
    else:
        print("CHANGE 2 FAILED: Could not find paragraph containing the target text.")

    # ------------------------------------------------------------------
    # CHANGE 3 - Paras 33-36 (Section 1.3.2 bullets)
    # Replace 4 old bullets with 5 new bullets
    # ------------------------------------------------------------------
    bullet_old_texts = [
        "Build a website and Chrome extension.",
        "Scrape and enrich Amazon product data",
        "Train an XGBoost",
        "Return an A+\u2192F grade with confidence",
    ]
    # Also try alternate substrings in case encoding differs
    bullet_old_texts_alt = [
        "Build a website and Chrome extension.",
        "Scrape and enrich Amazon product data",
        "Train an XGBoost",
        "Return an A+",
    ]

    new_bullet_texts = [
        ("Implement a three-tier material detection system capable of handling structured spec-table data, free-text percentage "
         "compositions, and title-keyword inference, falling back gracefully when source data is absent."),
        ("Train a multi-class XGBoost eco-grade classifier (A+\u2013F) on a DEFRA-derived synthetic dataset, achieving at least 90% "
         "cross-validated macro F1, and wrap it with post-hoc isotonic calibration and conformal prediction to provide statistically "
         "grounded uncertainty estimates."),
        ("Deploy a browser extension that injects eco-grade overlays into live Amazon product pages in real time without requiring "
         "the user to leave the purchase flow."),
        ("Provide per-prediction explainability through SHAP feature contributions and a six-stage LCA decomposition, so users "
         "understand what drives the grade rather than receiving an opaque score."),
        ("Evaluate the system honestly against the original MoSCoW requirements, including an explicit acknowledgement of the "
         "limitations introduced by using a synthetic training corpus."),
    ]

    # Find the 4 bullet paragraphs
    bullet_paras = []
    for search_text in bullet_old_texts_alt:
        i, p = find_para_by_substring(doc, search_text)
        if p is not None:
            bullet_paras.append((i, p))
        else:
            print(f"  WARNING: Could not find bullet containing: '{search_text}'")

    if len(bullet_paras) == 4:
        # Sort by index to ensure correct order
        bullet_paras.sort(key=lambda x: x[0])
        bullet_style = bullet_paras[0][1].style.name

        # Replace text of first 4 bullets with new bullet texts 0-3
        for j, (bi, bp) in enumerate(bullet_paras):
            set_para_text(bp, new_bullet_texts[j])
            print(f"  Bullet {j+1} replaced at para {bi}.")

        # Insert 5th bullet after the 4th bullet
        last_bullet_para = bullet_paras[3][1]
        new_para = insert_paragraph_after(last_bullet_para, new_bullet_texts[4], bullet_style, doc)
        if new_para is not None:
            print(f"  5th bullet inserted after para {bullet_paras[3][0]}.")
        else:
            print("  WARNING: 5th bullet insertion may have failed (paragraph object not found via search).")
            # Fallback: manually set text on the new element
            new_p_elem = last_bullet_para._element.getnext()
            if new_p_elem is not None:
                print("  Fallback: found new element via getnext(), setting text directly.")
                # Clear all 'w:r' (run) elements in new_p_elem
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                for r in new_p_elem.findall(f'{{{ns}}}r'):
                    new_p_elem.remove(r)
                # Add a new run
                r_elem = etree.SubElement(new_p_elem, f'{{{ns}}}r')
                t_elem = etree.SubElement(r_elem, f'{{{ns}}}t')
                t_elem.text = new_bullet_texts[4]

        print(f"CHANGE 3 applied: 4 bullet paragraphs replaced + 1 new bullet inserted.")
        changes_made += 1
    else:
        print(f"CHANGE 3 FAILED: Only found {len(bullet_paras)} of 4 expected bullet paragraphs.")
        # Print nearby paragraphs for debugging
        print("  Paragraphs around index 33-36:")
        for i in range(30, min(40, len(paragraphs))):
            print(f"    [{i}]: {paragraphs[i].text[:80]!r}")

    # ------------------------------------------------------------------
    # CHANGE 4 - Para 42 (Section 2.1 body text)
    # ------------------------------------------------------------------
    old_text_4 = "Bibliometric work on product disposal and sustainable consumption shows growing research attention on repair, reuse, recycling barriers and short product lifetimes."
    new_text_4 = ("Bibliometric work on product disposal and sustainable consumption shows growing research attention on repair, reuse, "
                  "recycling barriers, and short product lifetimes (Cruz-C\u00e1rdenas et al., 2022). For e-commerce, these themes imply "
                  "that interventions should target the purchase moment, when users can still choose alternatives. Importantly, the "
                  "reviewed literature concentrates on post-purchase disposal rather than pre-purchase choice, leaving the purchase "
                  "decision itself underserved as an intervention point \u2014 the gap this project occupies.")
    idx, para = find_para_by_substring(doc, old_text_4)
    if para is not None:
        set_para_text(para, new_text_4)
        print(f"CHANGE 4 applied: Para {idx} (Section 2.1) replaced.")
        changes_made += 1
    else:
        # Try shorter substring
        idx, para = find_para_by_substring(doc, "Bibliometric work on product disposal")
        if para is not None:
            set_para_text(para, new_text_4)
            print(f"CHANGE 4 applied: Para {idx} (Section 2.1) replaced (via shorter match).")
            changes_made += 1
        else:
            print("CHANGE 4 FAILED: Could not find paragraph.")

    # ------------------------------------------------------------------
    # CHANGE 5 - Para 55 (Section 2.2 body text)
    # ------------------------------------------------------------------
    old_text_5 = "LCA frameworks are the most rigorous way to locate lifecycle hotspots, but detailed inventories are rarely available for marketplace listings"
    new_text_5 = ("LCA frameworks are the most rigorous way to locate lifecycle hotspots, but detailed inventories are rarely available "
                  "for individual marketplace listings and are not usable in a seconds-to-decide shopping flow. Consumer-facing tools \u2014 "
                  "carbon calculators, eco-label databases \u2014 improve accessibility somewhat, but they sit outside the purchase journey "
                  "entirely, requiring a user to leave the product page, look up a score, and return. There is strong evidence that this "
                  "extra step is sufficient to prevent most users from acting on the information (Barreto et al., 2013). A key limitation "
                  "of existing tools is therefore the friction they impose on the user; this project\u2019s embedded browser extension "
                  "approach eliminates that friction by bringing the assessment to the product page rather than the reverse.")
    idx, para = find_para_by_substring(doc, old_text_5)
    if para is not None:
        set_para_text(para, new_text_5)
        print(f"CHANGE 5 applied: Para {idx} (Section 2.2) replaced.")
        changes_made += 1
    else:
        print("CHANGE 5 FAILED: Could not find paragraph.")

    # ------------------------------------------------------------------
    # CHANGE 6 - Para 57 (Section 2.3 body text) - APPEND to end
    # ------------------------------------------------------------------
    old_text_6 = "Behavioral research suggests that intention does not reliably translate into greener action"
    append_text_6 = (" The practical implication for system design is that framing environmental impact as a concrete, familiar-format "
                     "grade (A to F) at the moment of comparison is likely to be more effective than delayed reporting, separate "
                     "dashboards, or abstract CO\u2082 figures presented without context. Adoption barriers have similarly been identified "
                     "in technology-adjacent sustainability interventions, including smart energy systems (Sovacool et al., 2018), "
                     "reinforcing that usability and immediacy matter as much as information accuracy.")
    idx, para = find_para_by_substring(doc, old_text_6)
    if para is not None:
        # Append to end of paragraph
        current_text = para.text
        new_full_text = current_text + append_text_6
        set_para_text(para, new_full_text)
        print(f"CHANGE 6 applied: Para {idx} (Section 2.3) appended.")
        changes_made += 1
    else:
        print("CHANGE 6 FAILED: Could not find paragraph starting with 'Behavioral research'.")

    # ------------------------------------------------------------------
    # CHANGE 7 - Para 61 (Section 2.4 body text) - APPEND to end
    # ------------------------------------------------------------------
    old_text_7 = "Eco-feedback studies show that immediate, contextual feedback"
    append_text_7 = (" Most eco-feedback research has been conducted in home energy contexts; deployment within e-commerce at the point "
                     "of purchase remains largely unexplored in the literature, suggesting this project occupies a genuinely novel "
                     "application area.")
    idx, para = find_para_by_substring(doc, old_text_7)
    if para is not None:
        current_text = para.text
        new_full_text = current_text + append_text_7
        set_para_text(para, new_full_text)
        print(f"CHANGE 7 applied: Para {idx} (Section 2.4) appended.")
        changes_made += 1
    else:
        print("CHANGE 7 FAILED: Could not find paragraph starting with 'Eco-feedback studies'.")

    # ------------------------------------------------------------------
    # CHANGE 8 - Para 65 (Section 2.5 body text) - APPEND to end
    # ------------------------------------------------------------------
    old_text_8 = "Product pages rarely provide verified factory footprints"
    append_text_8 = (" The promise of AI for proxy-based scientific estimation has been noted more broadly (Venkatasubramanian, 2019); "
                     "however, none of the reviewed studies combine proxy-based estimation with conformal prediction to communicate "
                     "uncertainty bounds at a formally guaranteed coverage level \u2014 a gap this system directly addresses.")
    idx, para = find_para_by_substring(doc, old_text_8)
    if para is not None:
        current_text = para.text
        new_full_text = current_text + append_text_8
        set_para_text(para, new_full_text)
        print(f"CHANGE 8 applied: Para {idx} (Section 2.5) appended.")
        changes_made += 1
    else:
        print("CHANGE 8 FAILED: Could not find paragraph starting with 'Product pages rarely'.")

    # ------------------------------------------------------------------
    # CHANGE 9 - Para 82 (Section 2.9 Summary)
    # ------------------------------------------------------------------
    old_text_9 = "The literature supports point-of-purchase eco-feedback that is fast, interpretable and honest about uncertainty."
    new_text_9 = ("The literature supports point-of-purchase eco-feedback that is fast, interpretable, and honest about uncertainty. "
                  "Full LCA is rarely feasible for live marketplace products, so proxy-based estimation backed by open emission factors "
                  "and explainable ML is a reasonable practical compromise. These findings converge on three requirements not met by "
                  "any existing consumer tool identified in the review: (1) integration at the actual point of purchase rather than a "
                  "separate platform, (2) honest uncertainty communication so users understand the limits of estimates, and (3) "
                  "per-prediction feature attribution so users can understand what drives a grade rather than receiving an opaque score. "
                  "DSP Eco Tracker was designed to address all three, and each corresponds directly to a component in the architecture: "
                  "the browser extension addresses (1), conformal prediction addresses (2), and SHAP explanations address (3).")
    idx, para = find_para_by_substring(doc, old_text_9)
    if para is not None:
        set_para_text(para, new_text_9)
        print(f"CHANGE 9 applied: Para {idx} (Section 2.9) replaced.")
        changes_made += 1
    else:
        # Try shorter match
        idx, para = find_para_by_substring(doc, "The literature supports point-of-purchase eco-feedback")
        if para is not None:
            set_para_text(para, new_text_9)
            print(f"CHANGE 9 applied: Para {idx} (Section 2.9) replaced (via shorter match).")
            changes_made += 1
        else:
            print("CHANGE 9 FAILED: Could not find paragraph.")

    # ------------------------------------------------------------------
    # CHANGE 10 - Para 121 (Section 3.2 placeholder)
    # ------------------------------------------------------------------
    old_text_10 = "\u2588 FIGURE 9 TO INSERT \u2588"
    new_text_10 = ("Figure 9: DSP Eco Tracker web interface showing eco-grade result for an Amazon product scan, with eco-grade badge, "
                   "CO\u2082 estimate, SHAP driver breakdown, and six-stage LCA panel visible (author screenshot).")
    idx, para = find_para_by_substring(doc, old_text_10)
    if para is not None:
        set_para_text(para, new_text_10)
        print(f"CHANGE 10 applied: Para {idx} (Section 3.2 placeholder) replaced.")
        changes_made += 1
    else:
        # Try alternate search
        idx, para = find_para_by_substring(doc, "FIGURE 9 TO INSERT")
        if para is not None:
            set_para_text(para, new_text_10)
            print(f"CHANGE 10 applied: Para {idx} (Section 3.2 placeholder) replaced (via alternate match).")
            changes_made += 1
        else:
            print("CHANGE 10 FAILED: Could not find placeholder paragraph.")
            # Debug: print paragraphs around 121
            for i in range(118, min(125, len(paragraphs))):
                print(f"    [{i}]: {paragraphs[i].text[:100]!r}")

    # ------------------------------------------------------------------
    # CHANGE 11 - Para 136 (Section 4.2)
    # ------------------------------------------------------------------
    old_text_11 = "SMOTE was applied after label re-derivation to address class imbalance."
    new_text_11 = ("SMOTE was applied after label re-derivation to address class imbalance. Grades A+ and F are substantially rarer in "
                   "the distribution of synthetic products than B, C, and D. XGBoost was trained with 300 estimators, a maximum tree "
                   "depth of 7, and a learning rate of 0.08. Five-fold cross-validation on the balanced dataset gave a mean accuracy "
                   "of 99.17% and mean macro F1 of 0.99. These figures should be interpreted carefully: because both the training "
                   "labels and the test labels were derived from the same DEFRA-based formula, the model is being evaluated against "
                   "the function it was trained to approximate. The accuracy reflects internal consistency rather than agreement with "
                   "independently verified LCA ground truth \u2014 a distinction discussed further in Section 5.2.")
    idx, para = find_para_by_substring(doc, old_text_11)
    if para is not None:
        set_para_text(para, new_text_11)
        print(f"CHANGE 11 applied: Para {idx} (Section 4.2) replaced.")
        changes_made += 1
    else:
        print("CHANGE 11 FAILED: Could not find paragraph.")

    # ------------------------------------------------------------------
    # CHANGE 12 - Para 139 (Figure 10 caption)
    # ------------------------------------------------------------------
    old_text_12 = "Figure 10: Confusion matrix for XGBoost eco-grade classifier on held-out test set (mean accuracy 86.6%, macro F1 0.84)"
    new_text_12 = ("Figure 10: Confusion matrix for XGBoost eco-grade classifier on held-out test set (test accuracy 99.28%, macro F1 "
                   "0.99). All misclassifications occur between adjacent grade boundaries.")
    idx, para = find_para_by_substring(doc, old_text_12)
    if para is not None:
        set_para_text(para, new_text_12)
        print(f"CHANGE 12 applied: Para {idx} (Figure 10 caption) replaced.")
        changes_made += 1
    else:
        # Try shorter match
        idx, para = find_para_by_substring(doc, "Figure 10: Confusion matrix for XGBoost")
        if para is not None:
            set_para_text(para, new_text_12)
            print(f"CHANGE 12 applied: Para {idx} (Figure 10 caption) replaced (via shorter match).")
            changes_made += 1
        else:
            print("CHANGE 12 FAILED: Could not find Figure 10 caption.")

    # ------------------------------------------------------------------
    # CHANGE 13 - Para 167 (Stray pytest command) - DELETE paragraph
    # ------------------------------------------------------------------
    old_text_13 = 'Run "cd backend && python -m pytest tests/test_app.py -v"'
    idx, para = find_para_by_substring(doc, old_text_13)
    if para is not None:
        # Clear all text from the paragraph (effectively emptying it)
        for run in para.runs:
            run.text = ''
        # Also try to remove the paragraph element from the document
        p_elem = para._element
        p_elem.getparent().remove(p_elem)
        print(f"CHANGE 13 applied: Para {idx} (stray pytest command) deleted.")
        changes_made += 1
    else:
        # Try alternate search
        idx, para = find_para_by_substring(doc, "cd backend && python -m pytest")
        if para is not None:
            p_elem = para._element
            p_elem.getparent().remove(p_elem)
            print(f"CHANGE 13 applied: Para {idx} (stray pytest command) deleted (via alternate match).")
            changes_made += 1
        else:
            print("CHANGE 13 FAILED: Could not find stray pytest command paragraph.")

    # ------------------------------------------------------------------
    # CHANGE 14 - Para 166 (Figure 13 caption) - fix trailing quote/space
    # ------------------------------------------------------------------
    old_text_14 = 'Figure 13: pytest output confirming all 142 unit tests pass across 16 test classes (author screenshot)." '
    new_text_14 = "Figure 13: pytest output confirming all 142 unit tests pass across 16 test classes (author screenshot)."
    idx, para = find_para_by_substring(doc, old_text_14)
    if para is not None:
        set_para_text(para, new_text_14)
        print(f"CHANGE 14 applied: Para {idx} (Figure 13 caption) fixed.")
        changes_made += 1
    else:
        # Try matching without trailing space/quote
        idx, para = find_para_by_substring(doc, 'Figure 13: pytest output confirming all 142 unit tests pass')
        if para is not None:
            current = para.text
            print(f"  Found Figure 13 caption at para {idx}: {current!r}")
            if current.endswith('." ') or current.endswith('."') or '" ' in current:
                set_para_text(para, new_text_14)
                print(f"CHANGE 14 applied: Para {idx} (Figure 13 caption) fixed.")
                changes_made += 1
            else:
                print(f"CHANGE 14: Para found but text doesn't match expected pattern. Current: {current!r}")
                # Apply it anyway since user specified this change
                set_para_text(para, new_text_14)
                print(f"CHANGE 14 applied: Para {idx} (Figure 13 caption) set to clean text.")
                changes_made += 1
        else:
            print("CHANGE 14 FAILED: Could not find Figure 13 caption.")

    # ------------------------------------------------------------------
    # Save output
    # ------------------------------------------------------------------
    print()
    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print(f"Saved successfully.")
    print()
    print(f"Total changes applied: {changes_made} / 14")

    if changes_made < 14:
        print()
        print("WARNING: Not all changes were applied. See FAILED messages above.")


if __name__ == "__main__":
    main()
