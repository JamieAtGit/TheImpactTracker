"""
Generate dissertation Word document for Jamie Young (22023338)
DSP — Environmental Change: An Overconsumption Problem Regarding User Awareness
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (2.5 cm all round, standard for UWE) ─────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(0.98)

# ── Core styles ───────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)

# ── Helper functions ──────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def para(text, bold_prefix=None, space_after=8):
    """Add a paragraph. If bold_prefix is set, that part is bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run2 = p.add_run(text)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.35)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def page_break():
    doc.add_page_break()

# ── Title page ────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(60)
t.paragraph_format.space_after  = Pt(12)
r = t.add_run("Environmental Change – An Overconsumption Problem\nRegarding User Awareness")
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(18)

for line in [
    "Digital Systems Project (UFCFXK-30-3)",
    "Jamie Liam Young  |  Student No. 22023338",
    "Supervisor: Neil Phillips",
    "University of the West of England, Bristol",
    "2025–2026",
]:
    tp = doc.add_paragraph(line)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(6)
    for run in tp.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. DESIGN
# ══════════════════════════════════════════════════════════════════════════════
heading("Design", 1)

heading("Methodology", 2)
para(
    "Given the exploratory and incremental nature of this project, an Agile development "
    "methodology was adopted. Unlike Waterfall, which demands a fully specified system before "
    "any implementation begins, Agile permits requirements to evolve as the system grows. This "
    "suited the project well: early prototypes of the Amazon scraper revealed that product data "
    "quality was far more variable than anticipated, which drove the need for tiered material "
    "detection and confidence scoring — features that were not part of the original specification. "
    "Each sprint produced a testable, deployable increment, allowing continuous validation against "
    "real Amazon product pages rather than hypothetical test cases."
)
para("MoSCoW prioritisation was used to distinguish core functionality from enhancements:")

bullet("Amazon product scraping (ASIN, weight, material, origin); CO₂ estimate output with eco grade A+–F; Flask REST API; React web interface; persistent MySQL database.", bold_prefix="Must Have: ")
bullet("Chrome extension for in-browser use; user authentication and scan history; multi-material composite detection; SHAP-based feature explanations; conformal prediction intervals.", bold_prefix="Should Have: ")
bullet("Life cycle assessment breakdown; personal carbon timeline; admin review queue; postcode-based UK delivery distance.", bold_prefix="Could Have: ")
bullet("Third-party verified product certifications; real-time competitor comparison; mobile native app.", bold_prefix="Won't Have (this iteration): ")

heading("System Architecture", 2)
para(
    "The system follows a layered pipeline architecture with four distinct tiers: data acquisition, "
    "estimation engine, persistence, and presentation."
)
para(
    "Data acquisition is handled by a requests-based Amazon scraper that extracts the product title, "
    "weight, material specification table, brand, and ASIN. Material detection operates at three "
    "confidence tiers: Tier 1 parses structured specification table fields directly; Tier 2 applies "
    "regex pattern matching across the full page text; Tier 3 uses a title-keyword lookup with a "
    "dictionary of approximately 80 material signals. Origin detection mirrors this hierarchy, "
    "combining a known brand-origin mapping (covering 200+ brands) with country name extraction "
    "from the product description.",
    bold_prefix="Data acquisition. "
)
para(
    "Two parallel pathways produce a CO₂ estimate: a rule-based formula and a trained XGBoost "
    "classifier. The rule-based path uses material CO₂ intensity values (sourced from ecoinvent v3.9 "
    "and DEFRA 2023) combined with weight and transport distance to compute an absolute kg CO₂e figure. "
    "The ML path predicts an eco grade (A+ to F) from an 8-feature vector encoding material, transport "
    "mode, recyclability, origin, log-weight, weight bin, and two interaction terms. The final output "
    "blends both signals: if the ML confidence exceeds 60% and data quality is high, the ML grade "
    "takes precedence; otherwise the rule-based estimate governs.",
    bold_prefix="The estimation engine. "
)
para(
    "Persistence is managed by a MySQL database via SQLAlchemy ORM with five tables: users, products, "
    "scraped_products, emission_calculations, and admin_reviews. The deliberate separation of "
    "scraped_products (product metadata) from emission_calculations (CO₂ predictions) reflects the "
    "fact that the same product, scanned by different users from different postcodes, produces "
    "different delivery-distance emissions. A single ScrapedProduct row may therefore have many "
    "associated EmissionCalculation rows. The scraped_products table caches results by ASIN to avoid "
    "repeated scraping; the emission_calculations table records both the ML and rule-based predictions "
    "alongside the final emission, calculation method, and a three-tier data quality signal (high, "
    "medium, low). This allows post-hoc analysis of where the two estimation paths diverge.",
    bold_prefix="Persistence. "
)
para(
    "Presentation has two surfaces: a React single-page application deployed on Netlify, and a Chrome "
    "extension that injects a persistent floating widget directly onto Amazon product pages. Both "
    "communicate with the Flask API deployed on Railway. Deploying the API separately from both "
    "frontends meant CORS configuration needed explicit thought; allowed origins are restricted to the "
    "two known deployment domains, and a _headers file sets a Content Security Policy that restricts "
    "script and connection sources to the same two domains.",
    bold_prefix="Presentation. "
)
para(
    "A six-stage life cycle assessment component was added to the frontend, computing emissions across "
    "raw material extraction, manufacturing, packaging, international shipping, UK regional distribution, "
    "and last-mile delivery to the user's postcode."
)

heading("Technology Choices", 2)
para(
    "Flask was chosen over Django because the project required a lightweight API without the overhead "
    "of a full framework's template engine and admin system. Flask's application factory pattern "
    "(create_app) made the testing configuration clean: the test suite instantiates the application "
    "with a testing configuration flag that swaps the production MySQL URI for an in-memory SQLite "
    "database, meaning schema and ORM relationship logic is exercised without any external dependency."
)
para(
    "XGBoost was selected for the ML classifier due to its robustness under class imbalance and its "
    "native handling of encoded categorical features. Gradient-boosted trees were preferred over a "
    "neural network approach partly for interpretability — SHAP values are straightforward to compute "
    "for tree-based models — and partly because the compact feature space (8 features) did not justify "
    "the additional complexity of a deep learning architecture."
)
para(
    "React was chosen for the frontend because its component model maps naturally to the product card, "
    "grade badge, LCA breakdown, and scan history elements that compose the UI. Tailwind CSS was used "
    "for styling to keep the design consistent without writing bespoke CSS files."
)
para(
    "Session-based authentication was chosen over JWT tokens. For a SPA context where the API and "
    "frontend share a CORS relationship, session cookies are simpler to revoke (by deleting the "
    "server-side session) and do not require the token-expiry management that JWTs demand. The "
    "trade-off is that sessions require server-side state, which would complicate horizontal scaling, "
    "but this was not a concern at the project's scale."
)

heading("Security Design", 2)
para(
    "Security was considered at the design stage rather than added retrospectively. The API "
    "rate-limits unauthenticated prediction requests to 10 per hour using flask-limiter, protecting "
    "against scraping abuse. Password hashing uses PBKDF2-SHA256 via werkzeug.security with a minimum "
    "policy enforced at the API layer (minimum 8 characters, at least one uppercase letter, one digit). "
    "The username 'admin' is reserved and cannot be registered through the public sign-up endpoint. "
    "CORS is restricted to the two known frontend origins. HTTP security headers — including "
    "X-Frame-Options: DENY, X-Content-Type-Options: nosniff, and a Content-Security-Policy directive "
    "that restricts script execution to same-origin sources — are set via a _headers file served by "
    "Netlify on the frontend and enforced at the API level through Flask response headers."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. IMPLEMENTATION AND TESTING
# ══════════════════════════════════════════════════════════════════════════════
heading("Implementation and Testing", 1)

heading("Scraper", 2)
para(
    "The scraper (backend/scrapers/amazon/requests_scraper.py) sends authenticated-looking HTTP "
    "requests using a pool of real browser user-agent strings. The response HTML is parsed with "
    "BeautifulSoup. Weight extraction uses a regex chain targeting the product detail table first, "
    "then falls back to pattern matching across the full page text. Amazon frequently formats weights "
    "inconsistently across product categories — values appear as '450 grams', '0.45 kg', or '1 lb' "
    "depending on the seller. The scraper normalises all values to kilograms, handling unit conversion "
    "for pounds and ounces."
)
para(
    "Where weight cannot be extracted from the page at all, the system falls back to a "
    "category-and-title-based default weight lookup. This covers approximately 40 product categories "
    "(from washing machines at 65 kg down to USB cables at 0.05 kg) based on typical product mass "
    "ranges. This fallback is flagged in the data quality signal so the uncertainty is visible to "
    "the user rather than silently absorbed into the estimate."
)
para(
    "Material detection proved the most iterative component. The initial implementation used only "
    "title-keyword matching, which misclassified products where the material only appeared in the "
    "specification table or embedded in a free-text description. The three-tier approach was developed "
    "in response to these failures: Tier 1 extracts from the structured spec table and carries the "
    "highest confidence weight; Tier 2 applies a comprehensive regex library across all page text; "
    "Tier 3 falls back to title heuristics. For multi-material products (e.g. shoes with a leather "
    "upper and rubber sole), the scraper captures all detected materials as a JSON array stored in "
    "the materials_json column, and the emission calculation blends material intensities proportionally."
)
para(
    "The scraper also maintains a brand-origin mapping for approximately 200 brands, built iteratively "
    "as scan data accumulated. When a new brand appeared frequently in the live dataset without a "
    "known origin, it was researched and added. This manual process is the primary reliability "
    "limitation of the origin detection subsystem."
)

heading("Machine Learning Pipeline", 2)
para(
    "The ML pipeline begins with a 50,000-row synthetic training dataset generated by applying the "
    "same DEFRA-derived CO₂ formula across a grid of material, transport, weight, origin, and "
    "recyclability combinations. This synthetic approach is the pipeline's core limitation — discussed "
    "in the Evaluation — but it ensures the model's decision boundaries are internally consistent "
    "with the rule-based formula from the outset."
)
para(
    "An important labelling inconsistency was identified during development: the original dataset's "
    "grade labels had been assigned using an earlier, looser threshold set, and were only 26% "
    "consistent with the production DEFRA thresholds. The retrain.py script addresses this by "
    "re-deriving all labels from CO₂ values using the current co2_to_grade function before any "
    "training occurs, ensuring labels and production thresholds remain in sync. The grade thresholds "
    "— A+ (≤0.05 kg CO₂e), A (≤0.15), B (≤0.40), C (≤1.00), D (≤2.50), E (≤5.00), F (>5.00) — "
    "are defined in a single shared function used by both the training pipeline and the production "
    "prediction endpoint, preventing divergence."
)
para(
    "SMOTE was applied after label re-derivation to address class imbalance. Grades A+ and F are "
    "substantially rarer in the distribution of synthetic products than B, C, and D. Training without "
    "oversampling caused the model to under-predict rare classes; SMOTE synthesises new minority-class "
    "examples in the feature space to produce a balanced training set before the 80/20 train/test split."
)
para(
    "XGBoost was trained with 300 estimators, a maximum tree depth of 7, and a learning rate of "
    "0.08. These hyperparameters were selected empirically by evaluating macro F1 score on the "
    "held-out test set, prioritising the balance of performance across all grade classes over raw "
    "accuracy on common grades. Five-fold cross-validation on the balanced dataset gave a mean "
    "accuracy of 86.6% and mean macro F1 of 0.84."
)
para(
    "Post-hoc isotonic calibration (CalibratedClassifierCV with method='isotonic') was applied to "
    "correct for the probability overconfidence common in gradient boosted trees. Calibration was "
    "validated using a reliability diagram — comparing the model's stated confidence against actual "
    "grade frequency in a held-out calibration set. The calibrated model's confidence outputs are "
    "then used directly in the API response: predictions below 60% confidence fall back to the "
    "rule-based estimate, and the UI surfaces the confidence score alongside each grade."
)
para(
    "Conformal prediction wraps the calibrated model to provide statistically grounded uncertainty "
    "intervals. The conformal wrapper computes a non-conformity score on a held-out calibration set "
    "and produces a prediction set at a 90% coverage level: the set contains all grades whose "
    "non-conformity score falls below the empirical quantile. This gives a formal marginal coverage "
    "guarantee — in expectation, 90% of true grades fall within the returned set — without making "
    "parametric assumptions about the underlying distribution."
)
para(
    "SHAP values are computed for each prediction to produce counterfactual explanations: 'if this "
    "product were made of aluminium instead of plastic, the estimated emissions would increase by "
    "X kg CO₂e.' These are displayed in both the web application and Chrome extension to support "
    "the system's transparency goal."
)

heading("Data Flywheel", 2)
para(
    "Every successfully scraped product is appended to ml/live_scraped.csv with its computed CO₂ "
    "values. When retrain.py is invoked, it merges this live data with the 50k base dataset and "
    "retrains the model on the full combined corpus. The pipeline is designed to be triggered "
    "manually when the live dataset has grown by approximately 10%, to avoid unnecessary compute "
    "cost. The flywheel means the model's coverage of real Amazon product types improves over time, "
    "with real product titles, weights, and materials supplementing the formula-derived base data."
)

heading("Life Cycle Assessment", 2)
para(
    "The LCA component (frontend/website/src/components/LifecycleAssessment.jsx) decomposes the "
    "total emission estimate into six stages: raw material extraction, manufacturing, packaging, "
    "international shipping, UK regional distribution, and last-mile delivery. Each stage's emission "
    "factor is derived from published sources — DEFRA 2023 for transport, IEA 2023 for manufacturing "
    "energy intensity by product category, WRAP 2023 for packaging and recyclability, and "
    "ecoinvent v3.9 for material extraction intensities."
)
para(
    "The UK distribution and last-mile stages use the user's postcode, mapped to one of four UK "
    "regions (England, Scotland, Wales, Northern Ireland), to estimate an approximate distance from "
    "a national distribution centre. The LCA breakdown is intentionally presented as approximate — "
    "each stage is labelled with a confidence tier — because several stages (particularly "
    "manufacturing energy) are modelled from category-level averages rather than product-specific "
    "data. This transparency is deliberate: the goal is to communicate the shape of a product's "
    "impact across its life cycle, not to imply false precision in each individual figure."
)
para(
    "Recyclability data used across the LCA and the rule-based calculator draws from WRAP 2022/23 "
    "figures for the UK context specifically, rather than optimistic global averages. UK plastic "
    "recycling is coded at 12% (RECOUP 2022) rather than a commonly cited global average of 20%, "
    "and UK textile recycling at 15% — both substantially lower than figures often quoted in "
    "sustainability communications."
)

heading("Chrome Extension", 2)
para(
    "The Chrome extension (frontend/extension/src/overlay.js) injects a persistent floating widget "
    "into the browser DOM when activated on an Amazon product page. The widget is implemented as a "
    "content script that appends a styled div element to the page body, checking for an existing "
    "instance before injecting (via a DOM element ID check) to prevent duplicate overlays. User "
    "interaction state (minimised/expanded, last analysis results) is persisted using the "
    "chrome.storage.local API keyed by the current URL, so reopening a previously analysed product "
    "page restores the previous result without requiring a new API call."
)
para(
    "The extension's URL auto-fill button reads window.location.href and populates the URL input "
    "field, reducing friction when the user is already viewing the product they want to analyse. "
    "Communication with the Flask API follows the same request path as the web application, with "
    "the same CORS restrictions applied — the extension's manifest.json declares the Railway API "
    "URL as a permitted host to allow cross-origin requests from the extension context."
)

heading("API and Backend", 2)
para(
    "The Flask API exposes approximately 25 endpoints covering prediction, user authentication, "
    "scan history, personal statistics, admin review, model metrics, and health checking. Rate "
    "limiting restricts unauthenticated prediction requests to 10 per hour. Flask-Migrate manages "
    "schema migrations, meaning database schema changes are versioned and applied without data loss. "
    "The health endpoint reports both server status and whether the ML model is loaded in memory, "
    "providing an observable signal for deployment monitoring."
)
para(
    "The prediction endpoint is the most complex: it receives a URL and optional postcode, triggers "
    "the scraper, classifies material and origin, runs both the rule-based and ML estimation paths, "
    "stores the result, and returns a combined JSON response containing the CO₂ estimate, eco grade, "
    "ML confidence, conformal prediction set, SHAP values, data quality signal, and the full LCA "
    "stage breakdown. A threading lock prevents race conditions when two simultaneous requests "
    "attempt to lazy-load the ML model from disk — a concrete concurrency issue encountered "
    "during load testing."
)

heading("Testing", 2)
para(
    "The test suite (backend/tests/test_app.py) contains 142 passing tests organised across 16 "
    "test classes. The full suite runs against an in-memory SQLite database via the Flask test "
    "client, avoiding any dependency on the production MySQL instance. All 142 tests pass in "
    "approximately 12 seconds on a standard development machine."
)
para("Test class coverage includes:")

test_bullets = [
    ("Eco-grade thresholds (11 parametrised boundary tests): ", "verifies CO₂ values at and around each grade boundary, including exact boundary cases (e.g. exactly 5.00 kg must return E not F), and that grades are monotonically ordered as CO₂ increases."),
    ("CO₂ formula (5 tests): ", "validates correct absolute values for specific material/transport combinations, correct scaling with weight, and the expected ordering between transport modes and materials."),
    ("Authentication (10 tests): ", "covers sign-up validation (short passwords, missing uppercase, missing digit, duplicate usernames, reserved username 'admin'), login success and failure paths, and session clearing on logout."),
    ("Material detection (two test classes): ", "verifies that smart_guess_material correctly identifies common materials from product titles, and that detect_material handles both structured spec table text and free-text descriptions."),
    ("Category detection: ", "verifies that detect_category_from_title correctly categorises product titles into the expected category strings used by the weight fallback and LCA stage selection."),
    ("Weight extraction: ", "tests the regex weight parser across multiple formats including grams, kilograms, pounds, ounces, and multi-value ranges."),
    ("ML model (tests 11–13): ", "loads the calibrated model from disk, checks output dimensions (7 grade classes), verifies all confidence values lie within [0, 1], and checks that every returned grade label is a member of the valid set."),
    ("Conformal prediction: ", "checks the configuration structure and that coverage levels are within the valid range."),
    ("Data quality aggregation: ", "verifies that combined origin and material confidence tiers correctly produce the expected high, medium, or low data quality signal."),
    ("Postcode mapping: ", "verifies that representative postcodes for Wales, Scotland, Northern Ireland, and England map to their correct UK regions."),
    ("Cache TTL: ", "checks the 30-day freshness window logic — results older than 30 days trigger a fresh scrape rather than serving the cached calculation."),
]
for bold, rest in test_bullets:
    bullet(rest, bold_prefix=bold)

para(
    "No mocking of the database layer is used. The test fixture initialises a real SQLAlchemy "
    "session against in-memory SQLite, so schema relationships, foreign key constraints, and ORM "
    "queries are all exercised. This decision was deliberate: mocking the database would allow "
    "ORM-level bugs to pass silently, surfacing only in production."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROJECT EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading("Project Evaluation", 1)

heading("Against Requirements", 2)
para(
    "The MoSCoW requirements were substantially met. All must-have deliverables are implemented "
    "and deployed: the scraper reliably extracts product data from standard Amazon listing pages, "
    "the estimation engine returns a CO₂ figure and eco grade, and the Flask API serves both the "
    "web application and Chrome extension. The database correctly caches results by ASIN and "
    "records scan history per authenticated user."
)
para(
    "Among the should-haves, the Chrome extension, user authentication, multi-material detection, "
    "SHAP explanations, and conformal prediction intervals are all complete and deployed. The "
    "could-have features were selectively implemented: the LCA breakdown, carbon timeline, and "
    "postcode-based delivery distance are all present; the admin review queue exists in the "
    "database schema and API but the admin frontend was deprioritised in favour of the ML pipeline "
    "and test coverage."
)

heading("Limitations and Honest Appraisal", 2)
para(
    "The most significant limitation is the synthetic training data. Because labels are derived "
    "from the rule-based formula rather than independently validated measurements, the XGBoost "
    "model is effectively learning a regularised version of that formula. The ML path adds useful "
    "robustness — it handles unseen feature combinations more gracefully than a hard formula — "
    "but does not constitute an independent empirical check on emission estimates. A more rigorous "
    "system would use LCA database lookups at the product category level to generate labels, or "
    "employ expert-labelled real products as ground truth. The 86.6% held-out accuracy figure "
    "should be read in this light: the model is accurate relative to its own training labels, "
    "not relative to externally verified CO₂ values."
)
para(
    "Origin detection relies on a manually curated brand-origin map and country-name text matching. "
    "For products where the country of origin is not stated on the listing page and the brand is "
    "not in the map, the system defaults to China as the most probable origin for manufactured "
    "goods. This is a reasonable statistical heuristic, but it is not communicated to users in "
    "sufficient detail. The data quality signal (high/medium/low) partially addresses this, but "
    "a user who does not understand what the signal represents may take a low-confidence estimate "
    "at face value."
)
para(
    "The LCA component's manufacturing stage uses category-level energy intensity averages from "
    "IEA 2023 data. For products where manufacturing energy is dominant — electronics in "
    "particular — this averaging introduces substantial uncertainty. The LCA communicates "
    "confidence tiers per stage, but users may anchor on the total figure regardless."
)
para(
    "Amazon's scraping surface is inherently fragile. The requests-based scraper broke when Amazon "
    "updated its page layout on two occasions during the project. Each fix required inspecting the "
    "new DOM structure and updating the BeautifulSoup selectors. A production system would need "
    "continuous scraper monitoring and rapid-response maintenance to remain reliable."
)

heading("Testing Evaluation", 2)
para(
    "The decision to target 16 distinct test classes covering 142 cases was driven by the need to "
    "test pure logic functions in isolation without network or database dependencies. The eco-grade "
    "boundary tests and CO₂ formula tests run in milliseconds and provide high confidence in the "
    "numeric core of the system. Regression detection was the primary motivation: on two occasions "
    "during development, a refactor of the emission formula broke a boundary case that was caught "
    "immediately by the parametrised test suite."
)
para(
    "The test coverage is weakest for the scraper itself. Because the scraper makes outbound HTTP "
    "requests to Amazon, it cannot be exercised meaningfully without live requests or recorded-"
    "response fixtures. The test suite does not include scraper integration tests, meaning scraper "
    "regressions only become visible when a real scan fails. Introducing response-fixture-based "
    "scraper tests (using, for example, pytest-httpserver) would be a meaningful improvement."
)

heading("Reflexive Evaluation", 2)
para(
    "Looking back, the decision to use a synthetic training dataset was pragmatic given the project "
    "timeline but resulted in a fundamentally circular system: the ML model validates the rule-based "
    "formula rather than providing an independent check on it. If starting over, collecting even a "
    "few hundred real-world LCA figures from verified sources would have been more valuable than "
    "scaling the synthetic dataset to 50,000 rows, because it would allow the rule-based formula "
    "itself to be validated rather than merely replicated."
)
para(
    "The iterative approach to material detection worked well. This is a domain where edge cases "
    "are abundant and a specification-driven approach would have missed them: products with materials "
    "embedded in bullet-point lists, products whose spec table is dynamically loaded (and therefore "
    "absent from the static request response), and products where the 'material' table field contains "
    "a marketing phrase rather than a material name were all encountered only on real product pages. "
    "Each new failure mode was converted into a test case before the fix was written, gradually "
    "hardening the detection logic."
)
para(
    "The deployment split between Railway (backend) and Netlify (frontend) was driven by each "
    "platform's free tier, but introduced more CORS complexity than anticipated. Debugging CORS "
    "issues — particularly the interaction between preflight requests, session cookies, and the "
    "SameSite cookie attribute — consumed approximately two days of development time that a "
    "single-host deployment would not have required."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. FURTHER WORK AND CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading("Further Work and Conclusion", 1)

heading("Further Work", 2)
para(
    "The most impactful improvement would be replacing the synthetic training labels with a curated "
    "set of real LCA measurements. Even 1,000 products with verified CO₂ figures — sourced from "
    "manufacturer environmental product declarations, academic LCA databases, or the ecoinvent "
    "database directly — would provide an independent validation signal and likely surface categories "
    "where the DEFRA-formula approximation is systematically wrong. Electronics and clothing are two "
    "categories where real-world LCA values diverge substantially from simple material-weight "
    "formulae."
)
para(
    "A secondary priority would be automating the brand-origin map using a lightweight classifier "
    "trained on brand-origin pairs from a structured product database. The current hand-curated "
    "mapping covers 200 brands but requires manual maintenance as new brands enter the dataset."
)
para(
    "The Chrome extension currently requires the user to either copy-paste the product URL or click "
    "the auto-fill button. Native page-action activation — automatically populating the URL field "
    "and optionally triggering an estimate when the extension is opened on an Amazon product page "
    "— would further reduce friction. Browser extension Manifest V3 compatibility is also required "
    "for long-term Chrome Web Store distribution, as Manifest V2 support is being sunset."
)
para(
    "Adding a use-phase energy model for electronic products would substantially improve accuracy "
    "for a category where use-phase emissions often dominate total lifetime impact. A laptop used "
    "for four years at UK grid intensity consumes far more energy over its lifetime than was emitted "
    "in its manufacture — a nuance the current LCA does not capture."
)
para(
    "Finally, a longitudinal user study measuring whether eco grade exposure at point-of-purchase "
    "measurably changes purchasing decisions would test the core behavioural assumption underlying "
    "the entire project. Whether this holds for carbon labels specifically, given the psychological "
    "research on information overload and value-action gaps (Kollmuss and Agyeman, 2002), remains "
    "an open and important empirical question."
)

heading("Conclusion", 2)
para(
    "This project produced a full-stack system — web application, Chrome extension, Flask API, ML "
    "pipeline, and 142-test suite — that gives consumers a fast, transparent carbon footprint "
    "estimate for Amazon products at the point of purchase. The system blends a DEFRA-calibrated "
    "rule-based formula with a post-hoc isotonic calibrated XGBoost classifier, communicates "
    "uncertainty through conformal prediction intervals and tiered data quality signals, and "
    "decomposes estimates across a six-stage life cycle assessment."
)
para(
    "The overconsumption problem that motivated this project is not solved by any single tool, but "
    "the research literature consistently identifies information gaps as a significant barrier to "
    "pro-environmental behaviour (Thøgersen et al., 2010). By surfacing CO₂ estimates directly "
    "within the Amazon shopping interface at the moment of decision, the system attempts to close "
    "one of those gaps where it is most actionable. The data flywheel means the model's coverage "
    "of real product types improves as users scan products, and the conformal prediction framework "
    "ensures that growing confidence in predictions is earned statistically rather than assumed."
)
para(
    "The core technical challenge of this project — building a credible, honest, and explainable "
    "environmental impact estimate from sparse, inconsistent product listing data — proved harder "
    "than the literature's description of carbon labelling systems suggested. The result is a "
    "system that communicates what it does not know as clearly as what it does, which feels like "
    "the appropriate response to that difficulty."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
heading("Glossary", 1)

glossary_items = [
    ("ASIN", "Amazon Standard Identification Number. A 10-character alphanumeric identifier assigned to each product listed on Amazon."),
    ("Calibration (ML)", "The process of adjusting a model's output probabilities so that stated confidence levels correspond to empirical accuracy. A well-calibrated model that states 80% confidence is correct approximately 80% of the time."),
    ("Carbon footprint", "The total greenhouse gas emissions caused by an activity, product, or organisation, expressed in kilograms or tonnes of CO₂ equivalent (CO₂e)."),
    ("Conformal prediction", "A statistical framework that produces prediction sets with a formal coverage guarantee. For a 90% coverage level, the true label falls within the returned set 90% of the time across the marginal distribution."),
    ("CO₂e", "Carbon dioxide equivalent. A normalised unit expressing the global warming potential of a greenhouse gas relative to CO₂ over a 100-year horizon."),
    ("DEFRA", "Department for Environment, Food and Rural Affairs (UK). Publishes annual greenhouse gas conversion factors for converting activity data to CO₂e."),
    ("ecoinvent", "A Swiss life cycle inventory database providing background emission factors for materials, energy, and processes used in LCA studies."),
    ("Eco grade", "A letter grade (A+ to F) assigned to a product based on its estimated CO₂e emissions, using fixed DEFRA-derived thresholds."),
    ("Feature vector", "A numerical representation of a data instance used as input to a machine learning model. In this system, an 8-element array encoding material, transport mode, recyclability, origin, log-weight, weight bin, and two interaction terms."),
    ("Flask", "A lightweight Python web framework used to build the REST API backend."),
    ("LCA (Life Cycle Assessment)", "A method for quantifying the environmental impacts of a product across its full life cycle, from raw material extraction to end of life."),
    ("SHAP", "SHapley Additive exPlanations. A game-theoretic framework for interpreting machine learning predictions by attributing each feature's contribution to the model output."),
    ("SMOTE", "Synthetic Minority Over-sampling Technique. A method for addressing class imbalance in training data by generating synthetic examples in the feature space of under-represented classes."),
    ("SQLAlchemy", "A Python ORM used to interact with the MySQL database through Python class definitions rather than raw SQL queries."),
    ("XGBoost", "Extreme Gradient Boosting. An ensemble machine learning algorithm based on gradient-boosted decision trees, used here to classify products into eco grade categories."),
]

for term, definition in glossary_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    r1 = p.add_run(term + " — ")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(definition)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. TABLE OF ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Abbreviations", 1)

abbrevs = [
    ("API",     "Application Programming Interface"),
    ("ASIN",    "Amazon Standard Identification Number"),
    ("CORS",    "Cross-Origin Resource Sharing"),
    ("CO₂",     "Carbon Dioxide"),
    ("CO₂e",    "Carbon Dioxide Equivalent"),
    ("CSV",     "Comma-Separated Values"),
    ("DEFRA",   "Department for Environment, Food and Rural Affairs"),
    ("DOM",     "Document Object Model"),
    ("DSP",     "Digital Systems Project"),
    ("GHG",     "Greenhouse Gas"),
    ("HGV",     "Heavy Goods Vehicle"),
    ("IEA",     "International Energy Agency"),
    ("JSON",    "JavaScript Object Notation"),
    ("LCA",     "Life Cycle Assessment"),
    ("ML",      "Machine Learning"),
    ("MoSCoW",  "Must, Should, Could, Won't (prioritisation framework)"),
    ("NLP",     "Natural Language Processing"),
    ("ORM",     "Object Relational Mapper"),
    ("REST",    "Representational State Transfer"),
    ("SHAP",    "SHapley Additive exPlanations"),
    ("SMOTE",   "Synthetic Minority Over-sampling Technique"),
    ("SPA",     "Single Page Application"),
    ("SQL",     "Structured Query Language"),
    ("UWE",     "University of the West of England"),
    ("WEEE",    "Waste Electrical and Electronic Equipment"),
    ("WRAP",    "Waste and Resources Action Programme"),
    ("XGBoost", "Extreme Gradient Boosting"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Abbreviation"
hdr[1].text = "Expansion"
for cell in hdr:
    for para_cell in cell.paragraphs:
        for run in para_cell.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

for abbr, expansion in abbrevs:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = expansion
    for cell in row:
        for para_cell in cell.paragraphs:
            for run in para_cell.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/Users/jamie/Documents/University/ImpactTracker/Young_22023338_DSP_Dissertation_Sections.docx"
doc.save(out)
print(f"Saved: {out}")
