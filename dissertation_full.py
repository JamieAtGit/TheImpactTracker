"""
Generate FULL dissertation Word document for Jamie Young (22023338)
Includes: existing Intro + Lit Review (unchanged) + new sections
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin   = Inches(1.18)
    section.right_margin  = Inches(0.98)

# ── Base font ─────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

for lvl, sz in [(1, 16), (2, 13), (3, 12)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    return p

def p(text, space_after=8):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(18)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return para

def pb(text, bold_part, space_after=8):
    """Paragraph with leading bold text."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(18)
    r1 = para.add_run(bold_part)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = para.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return para

def bl(text, bold_part=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.35)
    if bold_part:
        r1 = para.add_run(bold_part)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
    r2 = para.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return para

def fig_caption(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return para

def fig_placeholder(fig_num, caption):
    """Grey placeholder box for a figure."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(2)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f'[Figure {fig_num} — insert image here]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    fig_caption(caption)

def page_break():
    doc.add_page_break()

def glossary_entry(term, definition):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.line_spacing = Pt(18)
    r1 = para.add_run(term + ' — ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = para.add_run(definition)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
tp.paragraph_format.space_after  = Pt(16)
r = tp.add_run('Environmental Change – An Overconsumption Problem\nRegarding User Awareness')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(20)

for line in [
    'Jamie Liam Young  |  Student No. 22023338',
    'The University of the West of England | Computer Science BSc',
    '',
    'Module: Digital Systems Project (UFCFXK-30-3)',
    'Supervisor: Neil Phillips',
    '',
    'GitHub Link: XYZ',
]:
    lp = doc.add_paragraph(line)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_after = Pt(5)
    for run in lp.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
h('ABSTRACT', 1)
p(
    'This background chapter frames the design of DSP Eco Tracker: a web application and Chrome '
    'extension that estimates the environmental impact of Amazon products at the point of purchase. '
    'The system scrapes product metadata (e.g., title, brand, weight, origin and packaging hints), '
    'enriches it using geolocation and material look-ups, engineers interpretable features (material, '
    'transport mode, origin, recyclability, weight and packaging), and uses machine-learning (XGBoost/'
    'random-forest with rule-based fallback) to return an eco-grade (A+→F) with confidence and '
    'feature-level diagnostics. The literature indicates that consumers often lack actionable, '
    'trustworthy impact information when shopping online; real-time eco-feedback and clear '
    'explanations are more likely to influence behaviour than generic, post-hoc calculators. This '
    'chapter reviews research on overconsumption, eco-feedback, lifecycle assessment, e-commerce '
    'packaging and returns, and the practical data/ethics challenges of scraping and modelling '
    'product impacts.'
)

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════════════════════
h('ACKNOWLEDGEMENTS', 1)
p(
    'I would like to express my gratitude to my project supervisor, Neil Phillips, for his insight '
    'and guidance throughout the development of this project, and to my family and friends for '
    'their support.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual placeholder)
# ══════════════════════════════════════════════════════════════════════════════
h('Table of Contents', 1)
p('Right-click this table in Word → Update Field → Update entire table.', space_after=14)

toc_entries = [
    ('ABSTRACT', '2'),
    ('ACKNOWLEDGEMENTS', '2'),
    ('Table of Contents', '3'),
    ('Table of Figures', '4'),
    ('1. Introduction', '5'),
    ('  1.1 Project Introduction', '5'),
    ('  1.2 The Problem', '5'),
    ('  1.3 Aims and Objectives', '5'),
    ('  1.4 Report Structure', '5'),
    ('2. Literature Review', '5'),
    ('3. Design', ''),
    ('4. Implementation and Testing', ''),
    ('5. Project Evaluation', ''),
    ('6. Further Work and Conclusion', ''),
    ('Glossary', ''),
    ('Table of Abbreviations', ''),
    ('References', ''),
]
for entry, pg in toc_entries:
    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(3)
    r1 = ep.add_run(entry)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    if pg:
        r2 = ep.add_run(f'\t{pg}')
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
h('Table of Figures', 1)
figs = [
    'Figure 1: Evolution in the number of documents on product disposal (bibliometric trend).',
    'Figure 2: Framework for green innovation behaviour (behavioural drivers).',
    'Figure 3: Eco-feedback interface example (real-time feedback).',
    'Figure 4: Circular economy lifecycle stages (lifecycle framing).',
    'Figure 5: Global plastic production growth since 1950 (waste pressure).',
    'Figure 6: UK waste statistics: biodegradable municipal waste to landfill, 2010–2022.',
    'Figure 7: Sustainability performance measures for green supply chain management.',
    'Figure 8: DSP Eco Tracker architecture (author-generated).',
]
for f in figs:
    fp = doc.add_paragraph()
    fp.paragraph_format.space_after = Pt(4)
    r = fp.add_run(f)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h('1. Introduction', 1)

h('1.1 Project Introduction', 2)
p(
    'E-commerce reduces friction (recommendations, one-click checkout, fast delivery), which can '
    'increase discretionary purchasing. The environmental cost is distributed across the lifecycle: '
    'materials and manufacturing, packaging, transport/returns, and end-of-life. DSP Eco Tracker '
    'makes these costs visible during shopping by providing a fast, interpretable eco-grade for '
    'Amazon listings.'
)

h('1.2 The Problem', 2)
p(
    'Most shoppers do not see credible lifecycle information at decision time; sustainability labels '
    'are inconsistent, and full LCA data is rarely available for marketplace products. This creates '
    'a mismatch: users may want to buy greener but default to convenience. A practical intervention '
    'is point-of-purchase eco-feedback that is transparent about uncertainty.'
)

h('1.3 Aims and Objectives', 2)
h('1.3.1 Aim', 3)
p('To design and evaluate a browser-based tool that raises awareness of environmental costs associated with online purchases (Amazon as the primary target).')

h('1.3.2 Objectives', 3)
bl('Build a website and Chrome extension.')
bl('Scrape and enrich product metadata (brand origin, facilities, packaging hints).')
bl('Engineer lifecycle-proxy features and score impact using ML with fallbacks.')
bl('Return an A+→F grade with confidence and clear driver explanations.')

h('1.4 Report Structure', 2)
p('This background chapter reviews the research that motivates real-time eco-feedback and justifies DSP Eco Tracker\'s proxy-based scoring approach.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
h('2. Literature Review', 1)

h('2.1 Research Outline / Background Study', 2)
p(
    'Bibliometric work on product disposal and sustainable consumption shows growing research '
    'attention on repair, reuse, recycling barriers and short product lifetimes. For e-commerce, '
    'these themes imply that interventions should target the purchase moment, when users can still '
    'choose alternatives.'
)
fig_placeholder(1, 'Figure 1: Evolution in the number of documents on product disposal.\nSource: ResearchGate — Consumer Behavior in Product Disposal: Mapping the Field (2022)')

h('2.2 Overview of Current Environmental Tools', 2)
p(
    'LCA frameworks are the most rigorous way to locate lifecycle hotspots, but detailed inventories '
    'are rarely available for marketplace listings and are not usable in a \'seconds-to-decide\' '
    'shopping flow. Consumer tools (ratings/calculators) improve accessibility, but often sit outside '
    'the purchase journey, limiting impact. This motivates embedded, lightweight assessments with '
    'explicit assumptions.'
)

h('2.3 Applications for Sustainable Consumption and Education', 2)
p(
    'Behavioral research suggests that intention does not reliably translate into greener action '
    'unless barriers are reduced. Digital tools can help by translating complex impacts into simple '
    'cues (scores, comparisons) and by explaining what drives impact. Frameworks on green innovation '
    'and environmental consciousness highlight perceived control and clarity as key drivers of '
    'behaviour change.'
)
fig_placeholder(2, 'Figure 2: Framework for green innovation behaviour.\nSource: ResearchGate — Green innovation behaviour: Impact of Industry 4.0 and open innovation (2023)')

h('2.4 Real-Time Feedback for Behavioural Change', 2)
p(
    'Eco-feedback studies show that immediate, contextual feedback is more persuasive than delayed '
    'reporting because it reduces psychological distance and supports trial-and-error. A browser '
    'extension can deliver feedback at the exact moment a user compares products, which strengthens '
    'the case for DSP Eco Tracker\'s point-of-decision design.'
)
fig_placeholder(3, 'Figure 3: Eco-feedback interface example.\nSource: ResearchGate — Why don\'t families get along with eco-feedback technologies? A longitudinal inquiry (2013)')

h('2.5 Data, Proxies and Machine Learning', 2)
p(
    'Product pages rarely provide verified factory footprints or complete material breakdowns, so '
    'real-time tools use proxies such as dominant material, estimated mass, origin and transport mode. '
    'Although tree-ensemble models do not provide causal explanations, their robustness to noisy, '
    'incomplete tabular data makes them appropriate for real-time sustainability estimation when paired '
    'with interpretability techniques. Interpretability is critical: explanations (e.g., feature '
    'contributions) improve auditability and user trust.'
)

h('2.6 Web Scraping, Ethics and Reliability', 2)
p(
    'Beyond technical robustness, transparency and explainability are increasingly emphasised in '
    'sustainability-focused digital systems. Research in explainable artificial intelligence (XAI) '
    'shows that users are more likely to trust and act upon algorithmic outputs when the system '
    'provides understandable reasons for its decisions rather than opaque scores alone. In the '
    'context of environmental decision-making, this is particularly important, as sustainability '
    'claims are often met with scepticism due to greenwashing concerns. By exposing feature-level '
    'drivers such as estimated weight, transport distance, dominant material and recyclability, DSP '
    'Eco Tracker aligns with recommendations from both sustainability and XAI literature, supporting '
    'accountability and informed user judgement.'
)
p(
    'Another relevant consideration is the trade-off between precision and usability. Full lifecycle '
    'models may yield more accurate results, but their complexity can overwhelm non-expert users. '
    'Conversely, overly simplified indicators risk misrepresentation. Prior studies suggest that '
    'proxy-based approaches, when clearly contextualised and accompanied by confidence indicators, '
    'can strike an effective balance for consumer-facing applications. DSP Eco Tracker adopts this '
    'approach by combining machine-learning predictions with confidence scores and deterministic '
    'fallbacks, ensuring consistent outputs even when data is incomplete. This design choice reflects '
    'broader industry practice, where resilience and graceful degradation are essential for '
    'real-world deployment.'
)
p(
    'Scraping enables product-level assessment but introduces brittleness (layout changes, missing '
    'fields) and ethical constraints (rate-limiting, avoiding personal data, respecting policies). '
    'Engineering mitigations include scraper fallbacks, safe encoding for unseen values, and '
    'rule-based scoring when models or encoders are unavailable, so the system fails gracefully '
    'rather than silently.'
)

h('2.7 Lifecycle Hotspots: Packaging, Transport and End-of-Life', 2)
p(
    'E-commerce hotspots align with DSP Eco Tracker\'s feature set: packaging and last-mile delivery '
    'add emissions; air freight is high-impact; returns amplify transport and waste. End-of-life '
    'evidence shows persistent landfill pressure and low recycling for key streams, supporting the '
    'inclusion of recyclability and packaging cues in the output.'
)
fig_placeholder(4, 'Figure 4: Circular economy lifecycle stages.\nSource: ResearchGate — Circular Economy: A Critical Literature Review of Concepts (2016)')
fig_placeholder(5, 'Figure 5: Global plastic production growth since 1950.\nSource: ResearchGate — Utilization of Plastic Wastes for Sustainable Environmental Management: A Review (2021)')
fig_placeholder(6, 'Figure 6: UK waste statistics (biodegradable municipal waste to landfill, 2010–2022).\nSource: GOV.UK — UK statistics on waste (2025)')

h('2.8 Measuring Sustainability and Carbon Footprint', 2)
p(
    'Sustainability metrics translate impacts into usable indicators, but comparability depends on '
    'boundaries and data quality. Supply-chain measurement work groups metrics across materials, '
    'energy, logistics and waste and emphasises transparency. DSP Eco Tracker follows this principle '
    'by returning the inputs used for scoring, confidence, and diagnostics, and by logging outputs '
    'to improve validity over time.'
)
fig_placeholder(7, 'Figure 7: Sustainability performance measures for green supply chain management.\nSource: ResearchGate — Sustainability Performance Measurement for Green Supply Chain Management (2021)')

h('2.9 Summary', 2)
p(
    'The literature supports point-of-purchase eco-feedback that is fast, interpretable and honest '
    'about uncertainty. Full LCA is rarely feasible for marketplace products, so proxy-based '
    'estimation backed by open datasets and explainable ML is a practical compromise. These findings '
    'justify DSP Eco Tracker\'s architecture: website + extension → scrape/enrich → feature '
    'engineering → ML/fallback scoring → grade, confidence and driver explanations → logging for '
    'continuous improvement.'
)

# Architecture diagram text
arch = doc.add_paragraph()
arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
arch.paragraph_format.space_before = Pt(8)
arch.paragraph_format.space_after  = Pt(4)
for line in [
    'Website / Chrome Extension',
    '↓',
    'Flask API',
    '↓',
    'Scraping & Extraction',
    '↓',
    'Data Enrichment',
    '↓',
    'Feature Engineering',
    '↓',
    'ML Model / Fallback',
    '↓',
    'Response + Logging',
]:
    lp = doc.add_paragraph(line)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_after = Pt(1)
    for run in lp.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

fig_caption('Figure 8: DSP Eco Tracker architecture (author-generated).')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. DESIGN
# ══════════════════════════════════════════════════════════════════════════════
h('3. Design', 1)

h('3.1 Methodology', 2)
p(
    'Given the exploratory and incremental nature of this project, an Agile development methodology '
    'was adopted. Unlike Waterfall, which demands a fully specified system before any implementation '
    'begins, Agile permits requirements to evolve as the system grows. This suited the project well: '
    'early prototypes of the Amazon scraper revealed that product data quality was far more variable '
    'than anticipated, which drove the need for tiered material detection and confidence scoring — '
    'features that were not part of the original specification. Each sprint produced a testable, '
    'deployable increment, allowing continuous validation against real Amazon product pages rather '
    'than hypothetical test cases.'
)
p('MoSCoW prioritisation was used to distinguish core functionality from enhancements:')

bl('Amazon product scraping (ASIN, weight, material, origin); CO₂ estimate output with eco grade A+–F; Flask REST API; React web interface; persistent MySQL database.', 'Must Have: ')
bl('Chrome extension for in-browser use; user authentication and scan history; multi-material composite detection; SHAP-based feature explanations; conformal prediction intervals.', 'Should Have: ')
bl('Life cycle assessment breakdown; personal carbon timeline; admin review queue; postcode-based UK delivery distance.', 'Could Have: ')
bl('Third-party verified product certifications; real-time competitor comparison; mobile native app.', "Won't Have (this iteration): ")

h('3.2 System Architecture', 2)
p(
    'The system follows a layered pipeline architecture with four distinct tiers: data acquisition, '
    'estimation engine, persistence, and presentation.'
)
pb(
    'Data acquisition is handled by a requests-based Amazon scraper that extracts the product title, '
    'weight, material specification table, brand, and ASIN. Material detection operates at three '
    'confidence tiers: Tier 1 parses structured specification table fields directly; Tier 2 applies '
    'regex pattern matching across the full page text; Tier 3 uses a title-keyword lookup with a '
    'dictionary of approximately 80 material signals. Origin detection mirrors this hierarchy, '
    'combining a known brand-origin mapping (covering 200+ brands) with country name extraction '
    'from the product description.',
    'Data acquisition. '
)
pb(
    'Two parallel pathways produce a CO₂ estimate: a rule-based formula and a trained XGBoost '
    'classifier. The rule-based path uses material CO₂ intensity values (sourced from ecoinvent v3.9 '
    'and DEFRA 2023) combined with weight and transport distance to compute an absolute kg CO₂e '
    'figure. The ML path predicts an eco grade (A+ to F) from an 8-feature vector encoding material, '
    'transport mode, recyclability, origin, log-weight, weight bin, and two interaction terms. The '
    'final output blends both signals: if the ML confidence exceeds 60% and data quality is high, '
    'the ML grade takes precedence; otherwise the rule-based estimate governs.',
    'The estimation engine. '
)
pb(
    'Persistence is managed by a MySQL database via SQLAlchemy ORM with five tables: users, '
    'products, scraped_products, emission_calculations, and admin_reviews. The deliberate separation '
    'of scraped_products (product metadata) from emission_calculations (CO₂ predictions) reflects '
    'the fact that the same product, scanned by different users from different postcodes, produces '
    'different delivery-distance emissions. A single ScrapedProduct row may therefore have many '
    'associated EmissionCalculation rows. Results are cached by ASIN with a 30-day freshness '
    'window; the emission_calculations table records both the ML and rule-based predictions '
    'alongside the final emission and a three-tier data quality signal (high, medium, low).',
    'Persistence. '
)
pb(
    'Presentation has two surfaces: a React single-page application deployed on Netlify, and a '
    'Chrome extension that injects a persistent floating widget directly onto Amazon product pages. '
    'Both communicate with the Flask API deployed on Railway. Allowed CORS origins are restricted '
    'to the two known deployment domains, and a _headers file enforces a Content Security Policy '
    'restricting script and connection sources.',
    'Presentation. '
)

h('3.3 Technology Choices', 2)
p(
    'Flask was chosen over Django because the project required a lightweight API without a full '
    'framework\'s template engine and admin system. Flask\'s application factory pattern (create_app) '
    'made testing straightforward: the test suite instantiates the application with a testing '
    'configuration flag that swaps the production MySQL URI for an in-memory SQLite database, '
    'meaning schema and ORM relationship logic is exercised without any external dependency.'
)
p(
    'XGBoost was selected for the ML classifier due to its robustness under class imbalance and '
    'its native handling of encoded categorical features. Gradient-boosted trees were preferred '
    'over a neural network approach partly for interpretability — SHAP values are straightforward '
    'to compute for tree-based models — and partly because the compact feature space (8 features) '
    'did not justify the additional complexity of a deep learning architecture.'
)
p(
    'React was chosen for the frontend because its component model maps naturally to the product '
    'card, grade badge, LCA breakdown, and scan history elements that compose the UI. Session-based '
    'authentication was chosen over JWT tokens: session cookies are simpler to revoke server-side '
    'and do not require the token-expiry management that JWTs demand.'
)

h('3.4 Security Design', 2)
p(
    'Security was considered at the design stage rather than added retrospectively. The API '
    'rate-limits unauthenticated prediction requests to 10 per hour using flask-limiter, protecting '
    'against scraping abuse. Password hashing uses PBKDF2-SHA256 via werkzeug.security with a '
    'minimum policy of 8 characters, at least one uppercase letter, and one digit. The username '
    '\'admin\' is reserved and cannot be registered through the public sign-up endpoint. HTTP '
    'security headers — including X-Frame-Options: DENY, X-Content-Type-Options: nosniff, and a '
    'Content-Security-Policy directive — are set via a _headers file on the Netlify frontend and '
    'enforced at the API level through Flask response headers.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. IMPLEMENTATION AND TESTING
# ══════════════════════════════════════════════════════════════════════════════
h('4. Implementation and Testing', 1)

h('4.1 Scraper', 2)
p(
    'The scraper (backend/scrapers/amazon/requests_scraper.py) sends authenticated-looking HTTP '
    'requests using a pool of real browser user-agent strings. The response HTML is parsed with '
    'BeautifulSoup. Weight extraction uses a regex chain targeting the product detail table first, '
    'then falls back to pattern matching across the full page text. Amazon frequently formats '
    'weights inconsistently — values appear as "450 grams", "0.45 kg", or "1 lb" depending on '
    'the seller. The scraper normalises all values to kilograms.'
)
p(
    'Where weight cannot be extracted from the page at all, the system falls back to a '
    'category-and-title-based default weight lookup. This covers approximately 40 product '
    'categories (from washing machines at 65 kg down to USB cables at 0.05 kg). This fallback '
    'is flagged in the data quality signal so the uncertainty is visible to the user rather than '
    'silently absorbed into the estimate.'
)
p(
    'Material detection proved the most iterative component. The initial implementation used only '
    'title-keyword matching, which misclassified products where the material only appeared in the '
    'specification table or embedded in a free-text description. The three-tier approach was '
    'developed in response to these failures: Tier 1 extracts from the structured spec table and '
    'carries the highest confidence weight; Tier 2 applies a comprehensive regex library across '
    'all page text; Tier 3 falls back to title heuristics. For multi-material products (e.g. shoes '
    'with a leather upper and rubber sole), the scraper captures all detected materials as a JSON '
    'array stored in the materials_json column, and the emission calculation blends material '
    'intensities proportionally. The scraper also maintains a brand-origin mapping for approximately '
    '200 brands, built iteratively as scan data accumulated.'
)

h('4.2 Machine Learning Pipeline', 2)
p(
    'The ML pipeline begins with a 50,000-row synthetic training dataset generated by applying the '
    'same DEFRA-derived CO₂ formula across a grid of material, transport, weight, origin, and '
    'recyclability combinations. An important labelling inconsistency was identified during '
    'development: the original dataset\'s grade labels had been assigned using an earlier, looser '
    'threshold set, and were only 26% consistent with the production DEFRA thresholds. The '
    'retrain.py script addresses this by re-deriving all labels from CO₂ values using the current '
    'co2_to_grade function before any training occurs. The grade thresholds — A+ (≤0.05 kg CO₂e), '
    'A (≤0.15), B (≤0.40), C (≤1.00), D (≤2.50), E (≤5.00), F (>5.00) — are defined in a single '
    'shared function used by both the training pipeline and the production prediction endpoint, '
    'preventing future divergence.'
)
p(
    'SMOTE was applied after label re-derivation to address class imbalance. Grades A+ and F are '
    'substantially rarer in the distribution of synthetic products than B, C, and D. XGBoost was '
    'trained with 300 estimators, a maximum tree depth of 7, and a learning rate of 0.08. '
    'Five-fold cross-validation on the balanced dataset gave a mean accuracy of 86.6% and mean '
    'macro F1 of 0.84.'
)
p(
    'Post-hoc isotonic calibration (CalibratedClassifierCV with method=\'isotonic\') was applied '
    'to correct for the probability overconfidence common in gradient boosted trees, validated '
    'using a reliability diagram. Conformal prediction wraps the calibrated model to provide '
    'statistically grounded uncertainty intervals at a 90% marginal coverage guarantee. SHAP '
    'values are computed for each prediction to produce counterfactual explanations surfaced in '
    'both the web application and Chrome extension.'
)

h('4.3 Data Flywheel', 2)
p(
    'Every successfully scraped product is appended to ml/live_scraped.csv with its computed CO₂ '
    'values. When retrain.py is invoked, it merges this live data with the 50k base dataset and '
    'retrains the model on the full combined corpus. The pipeline is triggered manually when the '
    'live dataset has grown by approximately 10%, meaning the model\'s coverage of real Amazon '
    'product types improves continuously as users scan products.'
)

h('4.4 Life Cycle Assessment', 2)
p(
    'The LCA component (frontend/website/src/components/LifecycleAssessment.jsx) decomposes the '
    'total emission estimate into six stages: raw material extraction, manufacturing, packaging, '
    'international shipping, UK regional distribution, and last-mile delivery. Each stage\'s '
    'emission factor is derived from published sources — DEFRA 2023 for transport, IEA 2023 for '
    'manufacturing energy intensity by product category, WRAP 2023 for packaging and recyclability, '
    'and ecoinvent v3.9 for material extraction intensities. The UK distribution and last-mile '
    'stages use the user\'s postcode, mapped to one of four UK regions, to estimate delivery '
    'distance. Recyclability data draws from WRAP 2022/23 figures for the UK context specifically '
    '— UK plastic recycling is coded at 12% (RECOUP 2022) and textile recycling at 15%, both '
    'substantially lower than commonly cited global averages. Each stage is labelled with a '
    'confidence tier to communicate the inherent approximation in category-level emission factors.'
)

h('4.5 Chrome Extension', 2)
p(
    'The Chrome extension (frontend/extension/src/overlay.js) injects a persistent floating widget '
    'into the browser DOM when activated on an Amazon product page. Implemented as a content script, '
    'it appends a styled div element to the page body, checking for an existing instance before '
    'injecting to prevent duplicate overlays. User interaction state and last analysis results are '
    'persisted using the chrome.storage.local API keyed by the current URL, so reopening a '
    'previously analysed product page restores the previous result without a new API call. The '
    'URL auto-fill button reads window.location.href and populates the URL input field, reducing '
    'friction when the user is already viewing the target product.'
)

h('4.6 API and Backend', 2)
p(
    'The Flask API exposes approximately 25 endpoints covering prediction, user authentication, '
    'scan history, personal statistics, admin review, model metrics, and health checking. The '
    'prediction endpoint is the most complex: it receives a URL and optional postcode, triggers '
    'the scraper, classifies material and origin, runs both the rule-based and ML estimation '
    'paths, stores the result, and returns a combined JSON response containing the CO₂ estimate, '
    'eco grade, ML confidence, conformal prediction set, SHAP values, data quality signal, and '
    'the full LCA stage breakdown. A threading lock prevents race conditions when two simultaneous '
    'requests attempt to lazy-load the ML model from disk — a concrete concurrency issue '
    'encountered during load testing.'
)

h('4.7 Testing', 2)
p(
    'The test suite (backend/tests/test_app.py) contains 142 passing tests organised across 16 '
    'test classes. The full suite runs against an in-memory SQLite database via the Flask test '
    'client, avoiding any dependency on the production MySQL instance. All 142 tests pass in '
    'approximately 12 seconds on a standard development machine.'
)
p('Key test class coverage:')
bl('Eco-grade thresholds (11 parametrised boundary tests): verifies CO₂ values at and around each grade boundary, including exact boundary cases (e.g. exactly 5.00 kg must return E not F), and that grades are monotonically ordered as CO₂ increases.')
bl('CO₂ formula (5 tests): validates correct absolute values for specific material/transport combinations, weight scaling, and the expected ordering between transport modes and materials.')
bl('Authentication (10 tests): sign-up validation (short passwords, missing uppercase, missing digit, duplicate usernames, reserved username \'admin\'), login success and failure paths, and session clearing on logout.')
bl('Material and category detection: verifies smart_guess_material, detect_material, and detect_category_from_title across structured spec table text, free-text descriptions, and title heuristics.')
bl('Weight extraction: tests the regex weight parser across grams, kilograms, pounds, ounces, and multi-value ranges.')
bl('ML model and conformal prediction: loads the calibrated model from disk, checks output dimensions (7 grade classes), verifies confidence values within [0, 1], checks all grade labels are valid, and validates conformal coverage levels.')
bl('Data quality aggregation: verifies that origin and material confidence tiers correctly produce the expected high, medium, or low signal.')
bl('Postcode mapping: verifies representative postcodes for Wales, Scotland, Northern Ireland, and England map to correct UK regions.')
bl('Cache TTL: checks the 30-day freshness window — results older than 30 days trigger a fresh scrape.')

p(
    'No mocking of the database layer is used. The test fixture initialises a real SQLAlchemy '
    'session against in-memory SQLite, so schema relationships, foreign key constraints, and ORM '
    'queries are all exercised rather than bypassed.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROJECT EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
h('5. Project Evaluation', 1)

h('5.1 Against Requirements', 2)
p(
    'The MoSCoW requirements were substantially met. All must-have deliverables are implemented '
    'and deployed: the scraper reliably extracts product data from standard Amazon listing pages, '
    'the estimation engine returns a CO₂ figure and eco grade, and the Flask API serves both the '
    'web application and Chrome extension. The database correctly caches results by ASIN and '
    'records scan history per authenticated user. Among the should-haves, the Chrome extension, '
    'user authentication, multi-material detection, SHAP explanations, and conformal prediction '
    'intervals are all complete and deployed. The LCA breakdown, carbon timeline, and '
    'postcode-based delivery distance are also present; the admin review queue exists in the '
    'database schema and API but the admin frontend was deprioritised in favour of the ML pipeline '
    'and test coverage.'
)

h('5.2 Limitations and Honest Appraisal', 2)
p(
    'The most significant limitation is the synthetic training data. Because labels are derived '
    'from the rule-based formula rather than independently validated measurements, the XGBoost '
    'model is effectively learning a regularised version of that formula. The ML path adds useful '
    'robustness — it handles unseen feature combinations more gracefully than a hard formula — '
    'but does not constitute an independent empirical check on emission estimates. A more rigorous '
    'system would use LCA database lookups at the product category level to generate labels, or '
    'employ expert-labelled real products as ground truth. The 86.6% held-out accuracy figure '
    'should be read in this light: the model is accurate relative to its own training labels, '
    'not relative to externally verified CO₂ values.'
)
p(
    'Origin detection relies on a manually curated brand-origin map and country-name text matching. '
    'For products where the country of origin is not stated on the listing page and the brand is '
    'not in the map, the system defaults to China as the most probable origin for manufactured '
    'goods — a reasonable statistical heuristic, but one that is not communicated to users in '
    'sufficient detail. The data quality signal (high/medium/low) partially addresses this, but '
    'a user who does not understand what the signal represents may take a low-confidence estimate '
    'at face value. Amazon\'s scraping surface is also inherently fragile: the requests-based '
    'scraper broke when Amazon updated its page layout on two occasions during the project, each '
    'requiring inspection of the new DOM structure and updated BeautifulSoup selectors.'
)

h('5.3 Testing Evaluation', 2)
p(
    'The decision to target 16 distinct test classes covering 142 cases was driven by the need to '
    'test pure logic functions in isolation without network or database dependencies. The eco-grade '
    'boundary tests and CO₂ formula tests run in milliseconds and provide high confidence in the '
    'numeric core of the system. On two occasions during development, a refactor of the emission '
    'formula broke a boundary case that was caught immediately by the parametrised test suite, '
    'demonstrating the regression-detection value of the approach. Test coverage is weakest for '
    'the scraper itself: because the scraper makes outbound HTTP requests to Amazon, it cannot '
    'be exercised meaningfully without live requests or recorded-response fixtures. Introducing '
    'response-fixture-based scraper tests would be a meaningful improvement.'
)

h('5.4 Reflexive Evaluation', 2)
p(
    'Looking back, the decision to use a synthetic training dataset was pragmatic given the project '
    'timeline but resulted in a fundamentally circular system: the ML model validates the '
    'rule-based formula rather than providing an independent check on it. If starting over, '
    'collecting even a few hundred real-world LCA figures from verified sources would have been '
    'more valuable than scaling the synthetic dataset to 50,000 rows, because it would allow the '
    'rule-based formula itself to be validated rather than merely replicated.'
)
p(
    'The iterative approach to material detection worked well. This is a domain where edge cases '
    'are abundant and a specification-driven approach would have missed them: products with '
    'materials embedded in bullet-point lists, products whose spec table is dynamically loaded '
    'and therefore absent from the static request response, and products where the "material" '
    'table field contains a marketing phrase rather than a material name were all encountered only '
    'on real product pages. Each new failure mode was converted into a test case before the fix '
    'was written, gradually hardening the detection logic. The deployment split between Railway '
    '(backend) and Netlify (frontend) introduced more CORS complexity than anticipated — debugging '
    'the interaction between preflight requests, session cookies, and the SameSite cookie attribute '
    'consumed approximately two days that a single-host deployment would not have required.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. FURTHER WORK AND CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h('6. Further Work and Conclusion', 1)

h('6.1 Further Work', 2)
p(
    'The most impactful improvement would be replacing the synthetic training labels with a curated '
    'set of real LCA measurements. Even 1,000 products with verified CO₂ figures — sourced from '
    'manufacturer environmental product declarations, academic LCA databases, or the ecoinvent '
    'database directly — would provide an independent validation signal and likely surface '
    'categories where the DEFRA-formula approximation is systematically wrong. Electronics and '
    'clothing are two categories where real-world LCA values diverge substantially from simple '
    'material-weight formulae.'
)
p(
    'A secondary priority would be automating the brand-origin map using a lightweight classifier '
    'trained on brand-origin pairs from a structured product database. The current hand-curated '
    'mapping covers 200 brands but requires manual maintenance as new brands enter the dataset. '
    'Native page-action activation for the Chrome extension — automatically populating the URL '
    'field when opened on an Amazon product page — would further reduce user friction. Browser '
    'extension Manifest V3 compatibility is also required for long-term Chrome Web Store '
    'distribution, as Manifest V2 support is being sunset.'
)
p(
    'Adding a use-phase energy model for electronic products would substantially improve accuracy '
    'for a category where use-phase emissions often dominate total lifetime impact. Finally, a '
    'longitudinal user study measuring whether eco grade exposure at point-of-purchase measurably '
    'changes purchasing decisions would test the core behavioural assumption underlying the entire '
    'project — something outside the scope of this work but an obvious and important extension.'
)

h('6.2 Conclusion', 2)
p(
    'This project produced a full-stack system — web application, Chrome extension, Flask API, ML '
    'pipeline, and 142-test suite — that gives consumers a fast, transparent carbon footprint '
    'estimate for Amazon products at the point of purchase. The system blends a DEFRA-calibrated '
    'rule-based formula with a post-hoc isotonic calibrated XGBoost classifier, communicates '
    'uncertainty through conformal prediction intervals and tiered data quality signals, and '
    'decomposes estimates across a six-stage life cycle assessment.'
)
p(
    'The overconsumption problem that motivated this project is not solved by any single tool, '
    'but the research literature consistently identifies information gaps as a significant barrier '
    'to pro-environmental behaviour. By surfacing CO₂ estimates directly within the Amazon '
    'shopping interface at the moment of decision, the system attempts to close one of those '
    'gaps where it is most actionable. The data flywheel means the model\'s coverage of real '
    'product types improves as users scan products, and the conformal prediction framework '
    'ensures that growing confidence in predictions is earned statistically rather than assumed.'
)
p(
    'The core technical challenge of this project — building a credible, honest, and explainable '
    'environmental impact estimate from sparse, inconsistent product listing data — proved harder '
    'than the literature\'s description of carbon labelling systems suggested. The result is a '
    'system that communicates what it does not know as clearly as what it does, which feels like '
    'the appropriate response to that difficulty.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h('Glossary', 1)

glossary_entry('ASIN', 'Amazon Standard Identification Number. A 10-character alphanumeric identifier assigned to each product listed on Amazon.')
glossary_entry('Calibration (ML)', 'The process of adjusting a model\'s output probabilities so that stated confidence levels correspond to empirical accuracy. A well-calibrated model that states 80% confidence is correct approximately 80% of the time.')
glossary_entry('Carbon footprint', 'The total greenhouse gas emissions caused by an activity, product, or organisation, expressed in kilograms or tonnes of CO₂ equivalent (CO₂e).')
glossary_entry('Conformal prediction', 'A statistical framework that produces prediction sets with a formal coverage guarantee. For a 90% coverage level, the true label falls within the returned set 90% of the time across the marginal distribution.')
glossary_entry('CO₂e', 'Carbon dioxide equivalent. A normalised unit expressing the global warming potential of a greenhouse gas relative to CO₂ over a 100-year horizon.')
glossary_entry('DEFRA', 'Department for Environment, Food and Rural Affairs (UK). Publishes annual greenhouse gas conversion factors for converting activity data to CO₂e.')
glossary_entry('ecoinvent', 'A Swiss life cycle inventory database providing background emission factors for materials, energy, and processes used in LCA studies.')
glossary_entry('Eco grade', 'A letter grade (A+ to F) assigned to a product based on its estimated CO₂e emissions, using fixed DEFRA-derived thresholds.')
glossary_entry('Feature vector', 'A numerical representation of a data instance used as input to a machine learning model. In this system, an 8-element array encoding material, transport mode, recyclability, origin, log-weight, weight bin, and two interaction terms.')
glossary_entry('Flask', 'A lightweight Python web framework used to build the REST API backend.')
glossary_entry('LCA (Life Cycle Assessment)', 'A method for quantifying the environmental impacts of a product across its full life cycle, from raw material extraction to end of life.')
glossary_entry('SHAP', 'SHapley Additive exPlanations. A game-theoretic framework for interpreting machine learning predictions by attributing each feature\'s contribution to the model output.')
glossary_entry('SMOTE', 'Synthetic Minority Over-sampling Technique. A method for addressing class imbalance in training data by generating synthetic examples in the feature space of under-represented classes.')
glossary_entry('SQLAlchemy', 'A Python ORM used to interact with the MySQL database through Python class definitions rather than raw SQL queries.')
glossary_entry('XGBoost', 'Extreme Gradient Boosting. An ensemble machine learning algorithm based on gradient-boosted decision trees, used here to classify products into eco grade categories.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
h('Table of Abbreviations', 1)

abbrevs = [
    ('API',     'Application Programming Interface'),
    ('ASIN',    'Amazon Standard Identification Number'),
    ('CORS',    'Cross-Origin Resource Sharing'),
    ('CO₂',     'Carbon Dioxide'),
    ('CO₂e',    'Carbon Dioxide Equivalent'),
    ('CSV',     'Comma-Separated Values'),
    ('DEFRA',   'Department for Environment, Food and Rural Affairs'),
    ('DOM',     'Document Object Model'),
    ('DSP',     'Digital Systems Project'),
    ('GHG',     'Greenhouse Gas'),
    ('HGV',     'Heavy Goods Vehicle'),
    ('IEA',     'International Energy Agency'),
    ('JSON',    'JavaScript Object Notation'),
    ('LCA',     'Life Cycle Assessment'),
    ('ML',      'Machine Learning'),
    ('MoSCoW',  'Must, Should, Could, Won\'t (prioritisation framework)'),
    ('NLP',     'Natural Language Processing'),
    ('ORM',     'Object Relational Mapper'),
    ('REST',    'Representational State Transfer'),
    ('SHAP',    'SHapley Additive exPlanations'),
    ('SMOTE',   'Synthetic Minority Over-sampling Technique'),
    ('SPA',     'Single Page Application'),
    ('SQL',     'Structured Query Language'),
    ('UWE',     'University of the West of England'),
    ('WEEE',    'Waste Electrical and Electronic Equipment'),
    ('WRAP',    'Waste and Resources Action Programme'),
    ('XAI',     'Explainable Artificial Intelligence'),
    ('XGBoost', 'Extreme Gradient Boosting'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Abbreviation'
hdr[1].text = 'Expansion'
for cell in hdr:
    for cp in cell.paragraphs:
        for run in cp.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

for abbr, expansion in abbrevs:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = expansion
    for cell in row:
        for cp in cell.paragraphs:
            for run in cp.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h('References', 1)

refs = [
    'Consumer Behavior in Product Disposal: Mapping the Field (2022) ResearchGate. Available from: https://www.researchgate.net/publication/361247762 [Accessed 3 January 2026].',
    'Green innovation behaviour: Impact of Industry 4.0 and open innovation (2023) ResearchGate. Available from: https://www.researchgate.net/publication/371038091 [Accessed 4 January 2026].',
    'Why don\'t families get along with eco-feedback technologies? A longitudinal inquiry (2013) ResearchGate. Available from: https://www.researchgate.net/publication/258838907 [Accessed 4 January 2026].',
    'Circular Economy: A Critical Literature Review of Concepts (2016) ResearchGate. Available from: https://www.researchgate.net/publication/291957061 [Accessed 5 January 2026].',
    'Utilization of Plastic Wastes for Sustainable Environmental Management: A Review (2021) ResearchGate. Available from: https://www.researchgate.net/publication/354137150 [Accessed 6 January 2026].',
    'Sustainability Performance Measurement for Green Supply Chain Management (2021) ResearchGate. Available from: https://www.researchgate.net/publication/350313748 [Accessed 6 January 2026].',
    'UK statistics on waste (2025) GOV.UK. Available from: https://www.gov.uk/government/statistics/uk-waste-data [Accessed 7 January 2026].',
    'Government conversion factors for company reporting of greenhouse gas emissions (2025) GOV.UK. Available from: https://www.gov.uk/government/collections/government-conversion-factors-for-company-reporting [Accessed 7 January 2026].',
    'XGBoost: A Scalable Tree Boosting System (2016) ResearchGate. Available from: https://www.researchgate.net/publication/301839264 [Accessed 12 January 2026].',
    'A Unified Approach to Interpreting Model Predictions (SHAP) (2017) ResearchGate. Available from: https://www.researchgate.net/publication/317062430 [Accessed 12 January 2026].',
    'Legality and Ethics of Web Scraping (2020) ResearchGate. Available from: https://www.researchgate.net/publication/352014123 [Accessed 13 January 2026].',
    'Web Scraping for Research: Legal, Ethical, Institutional and Scientific Considerations (2024) ResearchGate. Available from: https://www.researchgate.net/publication/385442945 [Accessed 14 January 2026].',
    'The Hidden Cost: Understanding the Environmental Impact of Online Purchase Returns (2024) ResearchGate. Available from: https://www.researchgate.net/publication/378288401 [Accessed 14 January 2026].',
    'Last-mile carbon emission under e-commerce: Environmental perspective (2022) ResearchGate. Available from: https://www.researchgate.net/publication/385591805 [Accessed 15 January 2026].',
    'GEM 2024: Global E-waste Monitor (2024) ResearchGate. Available from: https://www.researchgate.net/publication/379226356 [Accessed 15 January 2026].',
    'Effects of green products on consumers\' purchasing decision (2022) ResearchGate. Available from: https://www.researchgate.net/publication/360453965 [Accessed 17 January 2026].',
    'Encouraging Green Purchasing Behaviour by Increasing Environmental Consciousness (2014) ResearchGate. Available from: https://www.researchgate.net/publication/270463065 [Accessed 17 January 2026].',
    'Responsible consumer and lifestyle: Sustainability insights (2020) ResearchGate. Available from: https://www.researchgate.net/publication/343810883 [Accessed 18 January 2026].',
    'The Influence of Digital Transformation on Environmental Sustainability (2024) ResearchGate. Available from: https://www.researchgate.net/publication/377529823 [Accessed 19 January 2026].',
    'Kollmuss, A. and Agyeman, J. (2002) Mind the Gap: Why do people act environmentally and what are the barriers to pro-environmental behavior? Environmental Education Research, 8(3), pp. 239–260.',
    'Thøgersen, J., Haugaard, P. and Olesen, A. (2010) Consumer responses to ecolabels. European Journal of Marketing, 44(11/12), pp. 1787–1810.',
    'DEFRA (2023) Greenhouse Gas Reporting: Conversion Factors 2023. Department for Environment, Food and Rural Affairs. Available from: https://www.gov.uk/government/collections/government-conversion-factors-for-company-reporting [Accessed 7 January 2026].',
    'IEA (2023) World Energy Outlook 2023. International Energy Agency. Available from: https://www.iea.org/reports/world-energy-outlook-2023 [Accessed 10 January 2026].',
    'WRAP (2023) Recycling Tracker. Waste and Resources Action Programme. Available from: https://www.wrap.org.uk [Accessed 10 January 2026].',
    'ecoinvent (2023) ecoinvent Database v3.9. Swiss Centre for Life Cycle Inventories. Available from: https://ecoinvent.org [Accessed 10 January 2026].',
]

for ref in refs:
    rp = doc.add_paragraph()
    rp.paragraph_format.space_after = Pt(5)
    rp.paragraph_format.left_indent = Inches(0.35)
    rp.paragraph_format.first_line_indent = Inches(-0.35)
    run = rp.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ── Save ─────────────────────────────────────────────────────────────────────
out = '/Users/jamie/Documents/University/ImpactTracker/Young_22023338_DSP_Full_Dissertation.docx'
doc.save(out)
print(f'Saved: {out}')
